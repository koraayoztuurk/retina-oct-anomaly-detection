from __future__ import annotations

from pathlib import Path
import re

import matplotlib

matplotlib.use("Agg")
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt

from utils import save_text


TITLE_TR = "Normal Retina OCT Görüntülerinden Öğrenilen Konvolüsyonel Autoencoder ile Patolojik Örneklerin Yeniden Oluşturma Hatası Tabanlı Tespiti"
TITLE_EN = "Reconstruction-Error-Based Detection of Pathological Retinal OCT Images Using a Convolutional Autoencoder Trained on Normal Samples"
AUTHOR_ENTRIES = [
    {
        "name": "Koray Öztürk",
        "department": "Bilgisayar Mühendisliği",
        "institution": "Eskişehir Osmangazi Üniversitesi",
        "location": "Eskişehir/Türkiye",
        "email": "korayoztuurk@gmail.com",
    },
    {
        "name": "Emir Alp İlhan",
        "department": "Bilgisayar Mühendisliği",
        "institution": "Eskişehir Osmangazi Üniversitesi",
        "location": "Eskişehir/Türkiye",
        "email": "emiralpilhan@gmail.com",
    },
]

REFERENCE_ENTRIES = [
    {"key": "[1]", "citation": 'D. S. Kermany et al., "Identifying medical diagnoses and treatable diseases by image-based deep learning," Cell, vol. 172, no. 5, pp. 1122-1131.e9, 2018, doi: 10.1016/j.cell.2018.02.010.', "url": "https://doi.org/10.1016/j.cell.2018.02.010", "group": "OCT classification"},
    {"key": "[2]", "citation": 'D. S. Kermany, K. Zhang, and M. Goldbaum, "Large dataset of labeled optical coherence tomography (OCT) and chest X-ray images," Mendeley Data, ver. 3, 2018, doi: 10.17632/rscbjbr9sj.3.', "url": "https://doi.org/10.17632/rscbjbr9sj.3", "group": "Dataset"},
    {"key": "[3]", "citation": 'G. Litjens et al., "A survey on deep learning in medical image analysis," Med. Image Anal., vol. 42, pp. 60-88, 2017, doi: 10.1016/j.media.2017.07.005.', "url": "https://pubmed.ncbi.nlm.nih.gov/28778026/", "group": "Medical imaging survey"},
    {"key": "[4]", "citation": 'T. Schlegl, P. Seebock, S. M. Waldstein, U. Schmidt-Erfurth, and G. Langs, "Unsupervised anomaly detection with generative adversarial networks to guide marker discovery," arXiv:1703.05921, 2017, doi: 10.48550/arXiv.1703.05921.', "url": "https://arxiv.org/abs/1703.05921", "group": "Medical anomaly detection"},
    {"key": "[5]", "citation": 'S. Akcay, A. Atapour-Abarghouei, and T. P. Breckon, "GANomaly: Semi-supervised anomaly detection via adversarial training," in Proc. Asian Conf. Comput. Vis. (ACCV), pp. 622-637, 2018, doi: 10.1007/978-3-030-20893-6_39.', "url": "https://dblp.org/rec/conf/accv/AkcayAB18", "group": "Medical anomaly detection"},
    {"key": "[6]", "citation": 'V. Zavrtanik, M. Kristan, and D. Skocaj, "DRAEM - A discriminatively trained reconstruction embedding for surface anomaly detection," in Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV), pp. 8330-8339, 2021.', "url": "https://openaccess.thecvf.com/content/ICCV2021/html/Zavrtanik_DRAEM_-_A_Discriminatively_Trained_Reconstruction_Embedding_for_Surface_Anomaly_ICCV_2021_paper.html", "group": "Medical anomaly detection"},
    {"key": "[7]", "citation": 'P. Seebock et al., "Exploiting epistemic uncertainty of anatomy segmentation for anomaly detection in retinal OCT," IEEE Trans. Med. Imaging, vol. 39, no. 1, pp. 87-98, 2020, doi: 10.1109/TMI.2019.2919951.', "url": "https://dblp.org/rec/journals/tmi/SeebockOSWBKLS20", "group": "Retinal OCT anomaly detection"},
    {"key": "[8]", "citation": 'K. Zhou et al., "Proxy-bridged image reconstruction network for anomaly detection in medical images," IEEE Trans. Med. Imaging, vol. 41, no. 3, pp. 582-594, 2022, doi: 10.1109/TMI.2021.3118223.', "url": "https://research.nottingham.edu.cn/en/publications/proxy-bridged-image-reconstruction-network-for-anomaly-detection-", "group": "Medical anomaly detection"},
    {"key": "[9]", "citation": 'Y. Luo, Y. Ma, and Z. Yang, "Multi-resolution auto-encoder for anomaly detection of retinal imaging," Phys. Eng. Sci. Med., vol. 47, no. 2, pp. 517-529, 2024, doi: 10.1007/s13246-023-01381-x.', "url": "https://pubmed.ncbi.nlm.nih.gov/38285270/", "group": "Retinal OCT anomaly detection"},
    {"key": "[10]", "citation": 'J. Wang, W. Li, Y. Chen, et al., "Weakly supervised anomaly segmentation in retinal OCT images using an adversarial learning approach," Biomed. Opt. Express, vol. 12, no. 8, pp. 4713-4729, 2021, doi: 10.1364/BOE.426803.', "url": "https://pubmed.ncbi.nlm.nih.gov/34513220/", "group": "Retinal OCT anomaly detection"},
]


def format_metric(value: float) -> str:
    return f"{value:.4f}"


def frame_to_markdown(frame: pd.DataFrame, decimals: int | None = None) -> str:
    if decimals is not None:
        frame = frame.copy()
        numeric_columns = frame.select_dtypes(include=["number"]).columns
        frame[numeric_columns] = frame[numeric_columns].round(decimals)

    headers = list(frame.columns)
    header_row = "| " + " | ".join(headers) + " |"
    separator_row = "| " + " | ".join(["---"] * len(headers)) + " |"
    data_rows = ["| " + " | ".join(str(value) for value in row) + " |" for row in frame.itertuples(index=False, name=None)]
    return "\n".join([header_row, separator_row, *data_rows])


def threshold_table_markdown(threshold_table: pd.DataFrame) -> str:
    renamed = threshold_table.rename(
        columns={
            "percentile": "Persentil",
            "threshold": "Eşik",
            "accuracy": "Doğruluk",
            "precision": "Kesinlik",
            "recall": "Duyarlılık",
            "f1": "F1",
            "fpr": "FPR",
            "auroc": "AUROC",
        }
    )
    return frame_to_markdown(renamed, decimals=4)


def classwise_table_markdown(classwise_df: pd.DataFrame) -> str:
    renamed = classwise_df.rename(
        columns={
            "class_name": "Sınıf",
            "sample_count": "Örnek",
            "patient_count": "Hasta",
            "mean_reconstruction_error": "Ort. hata",
            "std_reconstruction_error": "Std. sapma",
        }
    )
    return frame_to_markdown(renamed, decimals=6)


def dataset_table_markdown(dataset_summary: pd.DataFrame) -> str:
    renamed = dataset_summary.replace(
        {
            "split_name": {
                "train": "Eğitim",
                "val": "Doğrulama",
                "test": "Test",
            }
        }
    ).rename(
        columns={
            "split_name": "Bölme",
            "class_name": "Sınıf",
            "image_count": "Görüntü",
            "patient_count": "Hasta",
        }
    )
    return frame_to_markdown(renamed)


def build_related_work_comparison_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"Çalışma": "Kermany et al. [1]", "Odak": "Denetimli OCT sınıflandırması", "Fark": "Patoloji etiketleri gerektirir; bizim yaklaşımımız yalnızca normal görüntülerle anomali tespiti yapar."},
            {"Çalışma": "AnoGAN [4]", "Odak": "GAN tabanlı anomali tespiti", "Fark": "Genel amaçlı anomali tespiti yaklaşımıdır; retinal OCT’ye özgü değildir."},
            {"Çalışma": "Seebock et al. [7]", "Odak": "Belirsizlik tabanlı OCT anomali tespiti", "Fark": "Doğrudan rekonstrüksiyon hatası yerine anatomi segmentasyonu belirsizliği kullanır."},
            {"Çalışma": "Luo et al. [9]", "Odak": "Çok çözünürlüklü retinal autoencoder", "Fark": "Daha gelişmiş retinal anomali modeli önerir; bizim çalışmamız ise daha sade ve tekrar üretilebilir bir temel model sunar."},
            {"Çalışma": "Wang et al. [10]", "Odak": "Zayıf denetimli retinal OCT anomali bölütleme", "Fark": "Anomali bölgelerinin konumunu vurgular; bizim yaklaşımımız görüntü düzeyinde puanlama yapan daha sade bir temel modeldir."},
            {"Çalışma": "Bu proje", "Odak": "Normal örneklerle OCT anomali puanlaması", "Fark": "Hasta düzeyinde doğrulama bölmesi ve gerçek OCT2017 verisi üzerinde persentil tabanlı eşik seçimi içerir."},
        ],
        columns=["Çalışma", "Odak", "Fark"],
    )


def build_experiment_setup_table(config: dict, dataset_summary: pd.DataFrame, history: dict) -> pd.DataFrame:
    train_count = int(dataset_summary.loc[(dataset_summary["split_name"] == "train") & (dataset_summary["class_name"] == "NORMAL"), "image_count"].iloc[0])
    val_count = int(dataset_summary.loc[(dataset_summary["split_name"] == "val") & (dataset_summary["class_name"] == "NORMAL"), "image_count"].iloc[0])
    test_count = int(dataset_summary.loc[dataset_summary["split_name"] == "test", "image_count"].sum())
    return pd.DataFrame(
        [
            {"Ayar": "Eğitim verisi", "Değer": f"{train_count} NORMAL görüntü"},
            {"Ayar": "Doğrulama verisi", "Değer": f"{val_count} NORMAL görüntü"},
            {"Ayar": "Test verisi", "Değer": f"{test_count} görüntü"},
            {"Ayar": "Giriş boyutu", "Değer": f"{config['image_size']}x{config['image_size']} gri seviye"},
            {"Ayar": "Latent boyut", "Değer": str(config["latent_dim"])},
            {"Ayar": "Optimizasyon", "Değer": f"Adam, lr={config['learning_rate']}"},
            {"Ayar": "Maks. epoch / patience", "Değer": f"{config['epochs']} / {config['early_stopping_patience']}"},
            {"Ayar": "Seçilen eşik", "Değer": f"p{config['default_percentile']}"},
            {"Ayar": "Eğitim süresi", "Değer": f"{history['training_time_sec'] / 60:.1f} dakika"},
        ]
    )


def references_markdown() -> str:
    lines = ["# Literature Notes", ""]
    grouped_entries: dict[str, list[dict]] = {}
    for entry in REFERENCE_ENTRIES:
        grouped_entries.setdefault(entry["group"], []).append(entry)

    for group_name, entries in grouped_entries.items():
        lines.extend([f"## {group_name}", ""])
        for entry in entries:
            lines.append(f"- {entry['key']} {entry['citation']}  ")
            lines.append(f"  Link: {entry['url']}")
    return "\n".join(lines) + "\n"


def create_pipeline_figure(save_path: Path) -> None:
    save_path.parent.mkdir(parents=True, exist_ok=True)
    fig, axis = plt.subplots(figsize=(12, 2.8))
    axis.set_xlim(0, 1)
    axis.set_ylim(0, 1)
    axis.axis("off")

    boxes = [
        (0.03, 0.28, 0.16, 0.42, "Yalnızca\nNORMAL eğitim"),
        (0.23, 0.28, 0.16, 0.42, "Yeniden boyutlandırma\n+ normalizasyon"),
        (0.43, 0.28, 0.16, 0.42, "Konvolüsyonel\nAutoencoder"),
        (0.63, 0.28, 0.16, 0.42, "Doğrulama\npersentilleri"),
        (0.83, 0.28, 0.14, 0.42, "Test anomali\npuanlaması"),
    ]

    for x, y, w, h, label in boxes:
        patch = plt.Rectangle((x, y), w, h, facecolor="#e9f2ff", edgecolor="#2d5f9a", linewidth=2)
        axis.add_patch(patch)
        axis.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=11)

    for start_x, end_x in [(0.19, 0.23), (0.39, 0.43), (0.59, 0.63), (0.79, 0.83)]:
        axis.annotate("", xy=(end_x, 0.49), xytext=(start_x, 0.49), arrowprops={"arrowstyle": "->", "lw": 2})

    axis.text(0.5, 0.88, "Retina OCT anomali tespit hattı", ha="center", va="center", fontsize=14, fontweight="bold")
    fig.tight_layout()
    fig.savefig(save_path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def create_results_overview_figure(output_root: Path, save_path: Path) -> None:
    candidates = [
        ("Eğitim kaybı", output_root / "figures" / "training_loss.png"),
        ("ROC eğrisi", output_root / "figures" / "roc_curve.png"),
        ("Hata dağılımı", output_root / "figures" / "test_error_distribution.png"),
        ("Rekonstrüksiyonlar", output_root / "reconstructions" / "reconstruction_examples.png"),
    ]
    available = [(title, path) for title, path in candidates if path.exists()]
    if not available:
        return

    rows = 2 if len(available) > 2 else 1
    cols = 2 if len(available) > 1 else 1
    fig, axes = plt.subplots(rows, cols, figsize=(11, 7))
    if rows == 1 and cols == 1:
        axes = [[axes]]
    elif rows == 1:
        axes = [axes]
    elif cols == 1:
        axes = [[axis] for axis in axes]

    flat_axes = [axis for row in axes for axis in row]
    for axis in flat_axes:
        axis.axis("off")

    for axis, (title, path) in zip(flat_axes, available):
        axis.imshow(mpimg.imread(path))
        axis.set_title(title, fontsize=11)
        axis.axis("off")

    fig.tight_layout()
    save_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(save_path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def build_markdown_report(
    config: dict,
    metrics: dict,
    thresholds: dict[int, float],
    threshold_table: pd.DataFrame,
    classwise_df: pd.DataFrame,
    dataset_summary: pd.DataFrame,
    history: dict,
    output_root: Path,
) -> str:
    related_work_table = build_related_work_comparison_table()
    setup_table = build_experiment_setup_table(config, dataset_summary, history)
    selected_percentile = config["default_percentile"]
    training_minutes = history["training_time_sec"] / 60

    return f"""# {TITLE_TR}

## Başlık

**İngilizce başlık:** {TITLE_EN}

## Özet

Bu çalışmada, retinal OCT görüntülerinde patolojik örnekleri etiketlenmiş patoloji sınıflarıyla doğrudan öğrenmek yerine, yalnızca normal örneklerden öğrenilen bir konvolüsyonel autoencoder ile yeniden oluşturma hatasına (reconstruction error) dayalı anomali tespiti yaklaşımı geliştirilmiştir. Kermany OCT2017 veri kümesindeki `train/NORMAL` görüntüleri hasta düzeyinde eğitim ve doğrulama alt kümelerine ayrılmış, model yalnızca normal anatominin dağılımını öğrenmiştir. Test aşamasında NORMAL, CNV, DME ve DRUSEN görüntüleri yeniden oluşturma hatasına göre puanlanmış ve doğrulama kümesindeki normal hata dağılımından elde edilen persentil eşikleriyle ikili karar üretilmiştir. Bu ara rapor sürümünde temel model 128x128 gri tonlamalı B-kesitleri üzerinde eğitilmiş, AUROC ana ölçüt olarak alınmış; doğruluk, kesinlik, duyarlılık, F1 ve yanlış pozitif oranı da raporlanmıştır. OCT2017 veri kümesi üzerinde yapılan deneylerde seçilen p{selected_percentile} eşiğinde AUROC {metrics['auroc']:.4f} ve F1 {metrics['f1']:.4f} elde edilmiştir.

## Abstract

This study investigates reconstruction-error-based anomaly detection for retinal OCT images using a convolutional autoencoder trained only on normal samples. Instead of learning a supervised pathology classifier, the model learns the appearance distribution of healthy retinal anatomy from the Kermany OCT2017 dataset. The training split uses only normal scans and a patient-level validation split is used to estimate operating thresholds from normal reconstruction-error percentiles. During testing, NORMAL, CNV, DME, and DRUSEN scans are scored according to reconstruction error, and anomaly decisions are produced with validation-derived thresholds. On the real OCT2017 experiment, the selected p{selected_percentile} operating point reaches AUROC {metrics['auroc']:.4f} and F1-score {metrics['f1']:.4f}. These findings indicate that a normal-only autoencoder can serve as a meaningful early baseline for retinal anomaly screening.

## Anahtar Kelimeler / Keywords

retinal OCT, anomali tespiti, autoencoder, yeniden oluşturma hatası, tıbbi görüntüleme, derin öğrenme

## 1. Giriş

Retinal hastalıkların erken tespiti, geri dönüşü olmayan görme kaybını azaltmak için kritik önemdedir. Optik koherens tomografi (OCT), retina tabakalarını yüksek çözünürlükte gösterebildiği için klinik pratikte sık kullanılan bir görüntüleme yöntemidir. Ancak OCT verisinin elle yorumlanması zaman alıcı olduğu gibi, geniş tarama programlarında yüksek uzman emeği gerektirir [1]. Son yıllarda derin öğrenme tabanlı denetimli modeller OCT sınıflandırmasında güçlü sonuçlar vermiş olsa da, bunlar genellikle her patoloji için etiketli veri gerektirir [1], [3]. Bu durum, daha önce görülmemiş veya yeterince temsil edilmeyen anomalilerin tespitini zorlaştırır.

Bu projede problem, normal anatominin öğrenilmesi ve ondan sapmaların yeniden oluşturma hatası ile yakalanması olarak ele alınmıştır. Ara rapor kapsamındaki amacımız, yalnızca normal retina OCT görüntüleri ile eğitilen bir konvolüsyonel autoencoder modelinin patolojik test görüntülerini anlamlı biçimde ayrıştırabildiğini gösteren, tekrar üretilebilir bir temel model kurmaktır. Bu çalışmada, hasta düzeyinde doğrulama, doğrulama dağılımından türetilen eşik seçimi ve gerçek OCT verisiyle uçtan uca çalışan bir deney hattını içeren tekrar üretilebilir bir temel sistem sunulmuştur.

## 2. İlgili Çalışmalar

OCT alanında derin öğrenme tabanlı hastalık sınıflandırması için en çok atıf alan çalışmalardan biri Kermany ve ark. tarafından sunulan Cell 2018 makalesidir [1]. Bu çalışma, aynı zamanda bu projede kullanılan halka açık OCT veri kümesinin temellerini de oluşturmaktadır [2]. Literatürde bunun devamında çok sayıda denetimli retinal hastalık tespit modeli önerilmiş ve OCT'nin otomatik analiz için uygunluğu güçlü biçimde ortaya konmuştur [3].

Anomali tespiti literatüründe ise normal veriyle eğitim yapıp anomalileri dağılım dışı örnekler olarak ele alan yeniden oluşturma temelli ve adversarial (çekişmeli) yöntemler ön plana çıkmıştır. AnoGAN [4] ve GANomaly [5] gibi yaklaşımlar normal dağılımı modelleme mantığını sistematikleştirmiştir. DRAEM [6] ve ProxyAno [8] ise yeniden oluşturma tabanlı yapıların daha ayırt edici hale gelmesine odaklanmıştır. Retinal OCT özelinde Seebock ve ark. [7], Luo ve ark. [9] ve Wang ve ark. [10] gibi çalışmalar bu alanın artık yalnızca genel amaçlı anomali tespiti değil, retina anatomisine özel çözümler de gerektirdiğini göstermektedir.

Kim ne yapmış ve bu proje neyi farklı yapıyor sorusunu daha açık göstermek için Tablo 1 verilmiştir.

Tablo 1. İlgili çalışmalar ve bu projeden farkları.

{frame_to_markdown(related_work_table)}

## 3. Yöntem

### 3.1 Veri kümesi ve bölme stratejisi

Çalışmada Kermany OCT2017 veri kümesinin `train` ve `test` klasörleri esas alınmıştır [2]. Eğitimde yalnızca `train/NORMAL` altındaki görüntüler kullanılmıştır. Doğrulama bölmesi görüntü düzeyinde değil hasta düzeyinde yapılmıştır; böylece aynı hastaya ait görüntüler eğitim ve doğrulama alt kümelerine aynı anda düşmemiştir. Hasta kimlikleri, veri kümesindeki dosya adlarında yer alan `hastalık-hastaID-görüntüNo` yapısından ayrıştırılmıştır. Test aşamasında `test/NORMAL`, `test/CNV`, `test/DME` ve `test/DRUSEN` görüntüleri birlikte değerlendirilmiş, NORMAL sınıfı 0 ve diğer tüm sınıflar anomali etiketi 1 olarak ele alınmıştır.

### 3.2 Ön işleme

Tüm görüntüler tek kanallı gri tonlamaya dönüştürülmüş, `128x128` boyutuna yeniden örneklenmiş ve `[0, 1]` aralığına normalize edilmiştir. Bu ara sürümde yoğun veri artırma uygulanmamıştır; amacımız önce sade ve tekrarlanabilir bir temel model kurmaktır.

### 3.3 Model mimarisi

Model, dört aşamalı bir konvolüsyonel autoencoder yapısından oluşmaktadır. Encoder kısmı `1->32->64->128->256` kanal geçişleri ve max-pooling adımlarıyla görüntüyü sıkıştırırken, ara gizil temsil `128` boyutlu bir vektöre indirgenmiştir. Decoder kısmı transpose convolution blokları ile görüntüyü tekrar `128x128` boyutuna taşımaktadır. Çıkış katmanında sigmoid kullanılarak normalize pikseller üzerinde yeniden oluşturma çıktısı üretilmiştir.

### 3.4 Eğitim ve eşikleme

Model `Adam` optimizer'ı ve `MSE` reconstruction loss ile eğitilmiştir. En fazla `40` epoch ve `8` patience değerli early stopping kullanılmıştır. Doğrulama aşamasında yalnızca normal örneklerin yeniden oluşturma hatası dağılımı incelenmiş; p95, p97 ve p99 eşikleri hesaplanmıştır. Ana operasyon noktası olarak p{selected_percentile} seçilmiştir. Böylece eşik seçiminde test verisi kullanılmamış ve data leakage engellenmiştir.

### 3.5 Değerlendirme ölçütleri ve deney kurulumu

Ana başarı ölçütü olarak AUROC seçilmiştir; çünkü anomali tespiti senaryosunda eşikten bağımsız ayrıştırma gücünü yansıtır. Bunun yanında doğruluk (accuracy), kesinlik (precision), duyarlılık (recall), F1 ve yanlış pozitif oranı (false positive rate, FPR) de raporlanmıştır. Kesinlik ve duyarlılık birlikte yorumlanmış, F1 ise dengeli operasyon noktası seçiminde kullanılmıştır. Gerçek deney koşusu yaklaşık {training_minutes:.1f} dakika sürmüş ve en iyi doğrulama sonucu {history['best_epoch']}. epoch'ta elde edilmiştir. Deney, NVIDIA GeForce RTX 4060 Laptop GPU içeren yerel bir PyTorch ortamında yürütülmüştür. Deney kurulumu Tablo 2'de özetlenmiştir.

Tablo 2. Deney kurulumu ve temel hiperparametreler.

{frame_to_markdown(setup_table)}

### 3.6 Sistem akışı

Önerilen iş akışı beş adımdan oluşmaktadır: normal verinin seçilmesi, ön işleme, autoencoder eğitimi, doğrulama hata dağılımından eşik seçimi ve testte anomali puanlaması. Şekil 1, önerilen sistemin iş akışını özetlemektedir.

## 4. Ara Sonuçlar

Bu ara raporda üretilen temel çıktılar; eğitim ve doğrulama kayıp grafiği, doğrulama yeniden oluşturma hatası histogramı, test hata dağılımı, ROC eğrisi, karışıklık matrisi ve örnek yeniden oluşturma-kalıntı görüntüleridir. Deney sonunda seçilen p{selected_percentile} eşiğinde elde edilen metrikler aşağıdaki gibidir:

| Metrik | Değer |
|---|---:|
| AUROC | {format_metric(metrics['auroc'])} |
| Doğruluk | {format_metric(metrics['accuracy'])} |
| Kesinlik | {format_metric(metrics['precision'])} |
| Duyarlılık | {format_metric(metrics['recall'])} |
| F1 | {format_metric(metrics['f1'])} |
| FPR | {format_metric(metrics['fpr'])} |
| En iyi epoch | {history['best_epoch']} |
| En iyi doğrulama kaybı | {history['best_val_loss']:.6f} |

Doğrulama persentil eşikleri:

{threshold_table_markdown(threshold_table)}

Sınıf bazlı yeniden oluşturma hatası özeti:

{classwise_table_markdown(classwise_df)}

Veri bölme özeti:

{dataset_table_markdown(dataset_summary)}

Sonuçlar yorum düzeyinde de anlamlıdır. p{selected_percentile} eşiği p97 ve p99'a göre daha yüksek duyarlılık ve F1 vermiştir; bu nedenle ara rapor için daha dengeli operasyon noktası olarak seçilmiştir. CNV ve DME sınıfları NORMAL görüntülerden belirgin şekilde ayrışırken, DRUSEN sınıfının hata dağılımı normale daha yakındır. Bu durum, bazı patolojilerin yeniden oluşturma tabanlı yaklaşımlarda diğerlerine göre daha zor ayrıştığını göstermektedir.

Şekil 2'de eğitim eğrisi, ROC performansı, hata dağılımı ve yeniden oluşturma örnekleri bir arada verilerek ara sonuçlar görsel olarak özetlenmiştir.

## 5. Tartışma

Temel model, görece basit olmasına rağmen normal anatomi dağılımını öğrenerek patolojik sınıfların yeniden oluşturma hatalarını yükseltebilmektedir. Bununla birlikte yeniden oluşturma tabanlı yöntemlerin iyi bilinen bir sınırı vardır: güçlü çözücü yapıları bazen anomalileri de fazla iyi yeniden üretebilir [4], [8]. Kermany veri kümesi görüntü düzeyinde etiketler içerir; bu nedenle yerel lezyon bölütlemesi için doğrudan piksel düzeyinde gerçek etiket bulunmamaktadır. Ayrıca eşik seçiminin kesinlik-duyarlılık dengesi üzerinde güçlü etkisi vardır. Bu nedenle tek bir metrik yerine persentil bazlı karşılaştırma tablosu korunmuştur.

Hesaplama maliyeti de göz ardı edilemez. Eğitim koşusu yerel ortamda uzun sayılabilecek bir sürede tamamlanmıştır ve bu durum veri yükleme ile ön işleme hattının da iyileştirme alanı olduğunu göstermektedir. Dolayısıyla mevcut sistem klinik kullanımdan ziyade araştırma ve erken tarama mantığında değerlendirilmelidir.

## 6. Gelecek Çalışmalar

Final aşamada ilk geliştirme ekseni, mimari seviyesinde daha güçlü yeniden oluşturma modellerinin denenmesi olacaktır. Standart autoencoder yerine VAE, skip-connection içeren daha derin yapılar veya bellek destekli yeniden oluşturma modelleri uygulanabilir. Buna ek olarak MSE yanında L1 ve SSIM tabanlı kayıplar denenerek yeniden oluşturma kalitesi ile anomali duyarlılığı arasındaki denge incelenecektir.

İkinci eksen, veri ve deney tasarımına odaklanacaktır. Daha yüksek giriş çözünürlüğü ile özellikle DRUSEN gibi daha zor ayrıştırılan sınıfların daha iyi temsil edilip edilmediği test edilecek; gizil boyut, mini-batch boyutu (batch size), eşik seçimi ve görüntü boyutu gibi hiperparametreler sistematik bir ablasyon çalışması ile karşılaştırılacaktır.

Üçüncü eksen, yorumlanabilirlik ve karşılaştırmalı değerlendirmedir. Kalıntı (residual) haritaları ve hata haritası görselleştirmeleri kullanılarak modelin hangi bölgelerde sapma ürettiği incelenecek; en iyi ve en kötü örnekler ayrıca tartışılacaktır. Bunun yanında veri yükleme hattı ve eğitim süresi optimize edilerek mevcut temel model en az bir geliştirilmiş varyantla AUROC, F1, duyarlılık ve FPR açısından karşılaştırılacaktır.

## 7. Sonuç

Bu ara rapor aşamasında, Kermany OCT verisi için hasta düzeyinde doğrulama kullanan, yalnızca normal görüntülerle eğitilen ve yeniden oluşturma hatası ile patolojik OCT örneklerini tespit eden tekrar üretilebilir bir temel sistem kurulmuştur. Gerçek veri üzerinde elde edilen AUROC {metrics['auroc']:.4f} ve F1 {metrics['f1']:.4f} değerleri, yaklaşımın umut verici olduğunu göstermektedir. Final aşamada hedef, bu temel modeli daha güçlü anomali tespiti yaklaşımlarıyla genişletmek ve sonuçları karşılaştırmalı deneylerle desteklemektir.

## Kaynaklar

""" + "\n".join(f"{entry['key']} {entry['citation']}" for entry in REFERENCE_ENTRIES) + "\n"


def _clear_document(doc: Document) -> None:
    body = doc._element.body
    for element in list(body):
        if element.tag.endswith("sectPr"):
            continue
        body.remove(element)


def _set_section_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_element = cols[0]
    else:
        cols_element = OxmlElement("w:cols")
        sect_pr.append(cols_element)
    cols_element.set(qn("w:num"), str(count))


def _set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_table_borders_black(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*(?:\.)?\s+", "", text).strip()


def _parse_markdown_sections(markdown_text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    for raw_line in markdown_text.strip().splitlines():
        if raw_line.startswith("# "):
            continue
        if raw_line.startswith("## "):
            if current_heading is not None:
                sections.append((current_heading, "\n".join(current_lines).strip()))
            current_heading = raw_line.replace("## ", "", 1).strip()
            current_lines = []
            continue
        if current_heading is not None:
            current_lines.append(raw_line)

    if current_heading is not None:
        sections.append((current_heading, "\n".join(current_lines).strip()))

    return sections


def _add_labeled_paragraph(doc: Document, label: str, content: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_run = paragraph.add_run(f"{label} - ")
    label_run.bold = True
    label_run.italic = True
    paragraph.add_run(" ".join(line.strip() for line in content.splitlines() if line.strip()))


def _add_author_block(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=len(AUTHOR_ENTRIES))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders_none(table)

    for column, author in enumerate(AUTHOR_ENTRIES):
        cell = table.rows[0].cells[column]
        lines = [
            author["name"],
            author["department"],
            author["institution"],
            author["location"],
            author["email"],
        ]
        for index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            if index == 0:
                run.bold = True


def _insert_inline_figure(doc: Document, caption: str, path: Path) -> None:
    if not path.exists():
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(2.9))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_body_sections(doc: Document, sections: list[tuple[str, str]], output_root: Path) -> None:
    figure_map = {
        "Yöntem": ("Şekil 1. Önerilen retina OCT anomali tespit boru hattı.", output_root / "figures" / "retina_oct_pipeline.png"),
        "Ara Sonuçlar": ("Şekil 2. Eğitim eğrisi, ROC performansı, hata dağılımı ve yeniden oluşturma örneklerinin özet görseli.", output_root / "figures" / "report_results_overview.png"),
    }

    for index, (heading, content) in enumerate(sections):
        normalized_heading = _strip_heading_number(heading)
        doc.add_heading(normalized_heading, level=1)
        if content.strip():
            _add_paragraphs(doc, content)

        if normalized_heading in figure_map:
            caption, path = figure_map[normalized_heading]
            _insert_inline_figure(doc, caption, path)


def _split_markdown_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    compact = "".join(cells).replace("-", "").replace(":", "").replace(" ", "")
    return compact == ""


def _add_markdown_table(doc: Document, lines: list[str]) -> None:
    parsed_rows = [_split_markdown_row(line) for line in lines if line.strip()]
    if len(parsed_rows) < 2:
        for line in lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return

    header = parsed_rows[0]
    data_rows = [row for row in parsed_rows[1:] if not _is_separator_row(row)]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    _set_table_borders_black(table)

    for column, cell_text in enumerate(header):
        cell = table.rows[0].cells[column]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for row_index, row_values in enumerate(data_rows, start=1):
        normalized = row_values + [""] * (len(header) - len(row_values))
        for column, cell_text in enumerate(normalized[: len(header)]):
            cell = table.rows[row_index].cells[column]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _add_paragraphs(doc: Document, text: str) -> None:
    lines = text.strip().splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip().replace("**", "")

        if not line:
            doc.add_paragraph("")
            index += 1
            continue

        if line.startswith("# "):
            index += 1
            continue

        if line.startswith("## "):
            doc.add_heading(_strip_heading_number(line.replace("## ", "", 1)), level=1)
            index += 1
            continue

        if line.startswith("### "):
            doc.add_heading(_strip_heading_number(line.replace("### ", "", 1)), level=2)
            index += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(doc, table_lines)
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(line[2:], style="List Bullet")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            index += 1
            continue

        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        index += 1


def build_docx_report(markdown_text: str, template_path: Path, output_path: Path, output_root: Path) -> None:
    if template_path.exists():
        doc = Document(template_path)
        _clear_document(doc)
    else:
        doc = Document()

    _set_section_columns(doc.sections[0], 1)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(TITLE_TR)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_paragraph.paragraph_format.space_after = Pt(6)

    english_title = doc.add_paragraph()
    english_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english_run = english_title.add_run(TITLE_EN)
    english_run.font.size = Pt(15)
    english_title.paragraph_format.space_after = Pt(10)

    _add_author_block(doc)
    doc.add_paragraph("")

    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    _set_section_columns(body_section, 2)

    sections = _parse_markdown_sections(markdown_text)
    section_map = {heading: content for heading, content in sections}
    body_sections = [
        (heading, content)
        for heading, content in sections
        if heading not in {"Başlık", "Özet", "Abstract", "Anahtar Kelimeler / Keywords"}
    ]

    if section_map.get("Özet"):
        _add_labeled_paragraph(doc, "Özet", section_map["Özet"])
    if section_map.get("Abstract"):
        _add_labeled_paragraph(doc, "Abstract", section_map["Abstract"])
    if section_map.get("Anahtar Kelimeler / Keywords"):
        _add_labeled_paragraph(doc, "Anahtar Kelimeler / Keywords", section_map["Anahtar Kelimeler / Keywords"])

    _add_body_sections(doc, body_sections, output_root)
    doc.save(output_path)


def build_report_assets(
    config: dict,
    metrics: dict,
    thresholds: dict[int, float],
    threshold_table: pd.DataFrame,
    classwise_df: pd.DataFrame,
    dataset_summary: pd.DataFrame,
    history: dict,
    output_root: Path,
    report_root: Path,
    template_path: Path,
) -> dict:
    report_root.mkdir(parents=True, exist_ok=True)

    create_pipeline_figure(output_root / "figures" / "retina_oct_pipeline.png")
    create_results_overview_figure(output_root, output_root / "figures" / "report_results_overview.png")

    markdown_report = build_markdown_report(
        config=config,
        metrics=metrics,
        thresholds=thresholds,
        threshold_table=threshold_table,
        classwise_df=classwise_df,
        dataset_summary=dataset_summary,
        history=history,
        output_root=output_root,
    )
    save_text(markdown_report, report_root / "ara_rapor_draft.md")
    save_text(references_markdown(), report_root / "literature_notes.md")
    build_docx_report(markdown_report, template_path, report_root / "ara_rapor_draft.docx", output_root)

    return {
        "title_tr": TITLE_TR,
        "title_en": TITLE_EN,
        "metrics": metrics,
        "thresholds": thresholds,
        "history": history,
        "threshold_table": threshold_table.round(6).to_dict(orient="records"),
        "classwise_summary": classwise_df.round(6).to_dict(orient="records"),
        "dataset_summary": dataset_summary.to_dict(orient="records"),
        "references": REFERENCE_ENTRIES,
        "generated_files": [
            str(report_root / "ara_rapor_draft.md"),
            str(report_root / "ara_rapor_draft.docx"),
            str(report_root / "literature_notes.md"),
            str(output_root / "figures" / "retina_oct_pipeline.png"),
            str(output_root / "figures" / "report_results_overview.png"),
        ],
    }
