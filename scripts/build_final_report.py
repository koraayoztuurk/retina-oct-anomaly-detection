from __future__ import annotations

import json
from pathlib import Path
import subprocess

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "report"
OUTPUT_DIR = PROJECT_ROOT / "outputs"
FINAL_FIGURE_DIR = OUTPUT_DIR / "final_report"
TEMPLATE_PATH = PROJECT_ROOT / "IEEE_Turkey_TUAC_Template_TR_2016_Final.docx"

TITLE_TR = (
    "Normal Retina OCT Görüntülerinden Öğrenilen Konvolüsyonel Autoencoder ile "
    "Patolojik Örneklerin Top-k Rekonstrüksiyon Hatası Tabanlı Tespiti"
)
TITLE_EN = (
    "Top-k Reconstruction-Error-Based Detection of Pathological Retinal OCT Images "
    "Using a Convolutional Autoencoder Trained on Normal Samples"
)

AUTHORS = [
    {
        "name": "Koray Öztürk",
        "department": "Bilgisayar Mühendisliği",
        "institution": "Eskişehir Osmangazi Üniversitesi",
        "location": "Eskişehir/Türkiye",
        "email": "korayoztuurk@gmail.com",
    },
    {
        "name": "Emir Alp İlhan",
        "department": "Bilgisayar Mühendisliği",
        "institution": "Eskişehir Osmangazi Üniversitesi",
        "location": "Eskişehir/Türkiye",
        "email": "emiralpilhan@gmail.com",
    },
]

FIGURE_EXPLANATION = (
    "Bulgular bölümündeki görseller yalnızca örnek görüntü göstermek için değil, tablolardaki sayısal sonuçların nasıl "
    "yorumlandığını desteklemek için kullanılmıştır. Şekil 2, final adayın hem genel model/score karşılaştırmasındaki "
    "yerini hem de sınıf bazlı tespit oranlarını tek bakışta özetler. Bu nedenle ana metrik tablolarından sonra verilmiştir."
    "\n\n"
    "Şekil 3, top-k residual score'un neyi yakaladığını nitel olarak göstermek için eklenmiştir. Her satırda orijinal görüntü, "
    "rekonstrüksiyon, residual haritası ve top-k residual overlay birlikte gösterilir. Bu görsel pixel-level segmentasyon iddiası "
    "taşımaz; yalnızca modelin hangi bölgelerde daha yüksek rekonstrüksiyon hatası ürettiğini açıklayıcı olarak gösterir."
    "\n\n"
    "Şekil 4, final sistemin en zayıf kaldığı DRUSEN false negative örneklerine ayrılmıştır. Bu örnekler, DRUSEN bulgularının "
    "bazı görüntülerde normal anatomiye yakın score üretebildiğini ve bu yüzden sınıf bazlı recall değerinin CNV ve DME'ye göre "
    "daha düşük kaldığını görsel olarak desteklemektedir."
)

REFERENCES = [
    '[1] D. S. Kermany et al., "Identifying medical diagnoses and treatable diseases by image-based deep learning," Cell, vol. 172, no. 5, pp. 1122-1131.e9, 2018, doi: 10.1016/j.cell.2018.02.010.',
    '[2] D. S. Kermany, K. Zhang, and M. Goldbaum, "Large dataset of labeled optical coherence tomography (OCT) and chest X-ray images," Mendeley Data, ver. 3, 2018, doi: 10.17632/rscbjbr9sj.3.',
    '[3] G. Litjens et al., "A survey on deep learning in medical image analysis," Med. Image Anal., vol. 42, pp. 60-88, 2017, doi: 10.1016/j.media.2017.07.005.',
    '[4] D. P. Kingma and M. Welling, "Auto-Encoding Variational Bayes," arXiv:1312.6114, 2013.',
    '[5] Z. Wang, A. C. Bovik, H. R. Sheikh, and E. P. Simoncelli, "Image quality assessment: From error visibility to structural similarity," IEEE Trans. Image Process., vol. 13, no. 4, pp. 600-612, 2004, doi: 10.1109/TIP.2003.819861.',
    '[6] T. Schlegl, P. Seebock, S. M. Waldstein, U. Schmidt-Erfurth, and G. Langs, "Unsupervised anomaly detection with generative adversarial networks to guide marker discovery," arXiv:1703.05921, 2017, doi: 10.48550/arXiv.1703.05921.',
    '[7] S. Akcay, A. Atapour-Abarghouei, and T. P. Breckon, "GANomaly: Semi-supervised anomaly detection via adversarial training," in Proc. Asian Conf. Comput. Vis. (ACCV), pp. 622-637, 2018, doi: 10.1007/978-3-030-20893-6_39.',
    '[8] V. Zavrtanik, M. Kristan, and D. Skocaj, "DRAEM - A discriminatively trained reconstruction embedding for surface anomaly detection," in Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV), pp. 8330-8339, 2021.',
    '[9] P. Seebock et al., "Exploiting epistemic uncertainty of anatomy segmentation for anomaly detection in retinal OCT," IEEE Trans. Med. Imaging, vol. 39, no. 1, pp. 87-98, 2020, doi: 10.1109/TMI.2019.2919951.',
    '[10] K. Zhou et al., "Proxy-bridged image reconstruction network for anomaly detection in medical images," IEEE Trans. Med. Imaging, vol. 41, no. 3, pp. 582-594, 2022, doi: 10.1109/TMI.2021.3118223.',
    '[11] Y. Luo, Y. Ma, and Z. Yang, "Multi-resolution auto-encoder for anomaly detection of retinal imaging," Phys. Eng. Sci. Med., vol. 47, no. 2, pp. 517-529, 2024, doi: 10.1007/s13246-023-01381-x.',
    '[12] J. Wang, W. Li, Y. Chen, et al., "Weakly supervised anomaly segmentation in retinal OCT images using an adversarial learning approach," Biomed. Opt. Express, vol. 12, no. 8, pp. 4713-4729, 2021, doi: 10.1364/BOE.426803.',
]


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def fmt(value: float, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def short_ci(ci_frame: pd.DataFrame, level: str, score_mode: str, metric: str) -> str:
    row = ci_frame[
        (ci_frame["level"] == level)
        & (ci_frame["score_mode"] == score_mode)
        & (ci_frame["metric"] == metric)
    ].iloc[0]
    return f"{fmt(row['point'])} [{fmt(row['ci_low'])}-{fmt(row['ci_high'])}]"


def markdown_table(frame: pd.DataFrame) -> str:
    headers = list(frame.columns)
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in frame.itertuples(index=False, name=None):
        lines.append("| " + " | ".join(str(value) for value in row) + " |")
    return "\n".join(lines)


def load_report_data() -> dict:
    ledger = pd.read_csv(OUTPUT_DIR / "experiment_ledger.csv")
    final_run = "ae_mse_l128_e60"
    score_root = OUTPUT_DIR / "score_ablation" / final_run
    return {
        "ledger": ledger,
        "config": load_json(OUTPUT_DIR / "experiments" / final_run / "metrics" / "run_config.json"),
        "history": load_json(OUTPUT_DIR / "experiments" / final_run / "metrics" / "training_history.json"),
        "best_image": load_json(score_root / "best_score.json"),
        "best_patient": load_json(score_root / "best_patient_score.json"),
        "classwise_score": pd.read_csv(score_root / "classwise_score_summary.csv"),
        "bootstrap": pd.read_csv(score_root / "bootstrap_confidence_intervals.csv"),
        "dataset": pd.read_csv(OUTPUT_DIR / "experiments" / final_run / "metrics" / "dataset_summary.csv"),
    }


def create_pipeline_figure(path: Path) -> None:
    FINAL_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8.2, 2.25))
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    steps = [
        ("NORMAL\ntrain", 0.03),
        ("Hasta bazlı\nval split", 0.22),
        ("AE / VAE\ntraining", 0.42),
        ("p95 eşik\nval normal", 0.62),
        ("Test score\n+ analiz", 0.81),
    ]
    for label, x in steps:
        box = plt.Rectangle((x, 0.33), 0.15, 0.34, facecolor="#edf4ff", edgecolor="#1f4e79", linewidth=1.5)
        ax.add_patch(box)
        ax.text(x + 0.075, 0.50, label, ha="center", va="center", fontsize=8.2)
    for x1, x2 in [(0.18, 0.22), (0.37, 0.42), (0.57, 0.62), (0.77, 0.81)]:
        ax.annotate("", xy=(x2, 0.50), xytext=(x1, 0.50), arrowprops={"arrowstyle": "->", "lw": 1.4})
    ax.text(0.5, 0.84, "Önerilen final deney hattı", ha="center", va="center", fontsize=10, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def create_results_figure(data: dict, path: Path) -> None:
    ledger = data["ledger"]
    final_row = ledger[
        (ledger["category"] == "score_ablation_image_level")
        & (ledger["run_id"] == "ae_mse_l128_e60")
        & (ledger["score_mode"] == "topk_mse_5")
    ].iloc[0]
    baseline = ledger[(ledger["category"] == "training_run") & (ledger["run_id"] == "ae_mse_l128")].iloc[0]
    vae = ledger[(ledger["category"] == "training_run") & (ledger["run_id"] == "vae_msekl_l128")].iloc[0]
    ssim = ledger[(ledger["category"] == "training_run") & (ledger["run_id"] == "ae_mse_ssim_l128")].iloc[0]

    classwise = data["classwise_score"]
    classwise = classwise[classwise["score_mode"] == "topk_mse_5"].copy()
    classwise["rate"] = classwise["detected_count"] / classwise["sample_count"]

    fig, axes = plt.subplots(1, 2, figsize=(8.2, 3.0))
    labels = ["AE-MSE\n40e", "VAE+KL", "MSE+SSIM", "AE top-k\n60e"]
    aurocs = [baseline["auroc"], vae["auroc"], ssim["auroc"], final_row["auroc"]]
    f1s = [baseline["f1"], vae["f1"], ssim["f1"], final_row["f1"]]
    x = range(len(labels))
    axes[0].bar([i - 0.18 for i in x], aurocs, width=0.36, label="AUROC", color="#2f6f9f")
    axes[0].bar([i + 0.18 for i in x], f1s, width=0.36, label="F1", color="#f28e2b")
    axes[0].set_xticks(list(x), labels, fontsize=7)
    axes[0].set_ylim(0.60, 1.00)
    axes[0].grid(axis="y", alpha=0.25)
    axes[0].legend(fontsize=7, loc="lower right")
    axes[0].set_title("Model/score karşılaştırması", fontsize=9)

    class_order = ["NORMAL", "CNV", "DME", "DRUSEN"]
    classwise = classwise.set_index("class_name").loc[class_order].reset_index()
    colors = ["#8f8f8f", "#4e79a7", "#59a14f", "#e15759"]
    axes[1].bar(classwise["class_name"], classwise["rate"], color=colors)
    axes[1].set_ylim(0, 1.05)
    axes[1].grid(axis="y", alpha=0.25)
    axes[1].set_title("Final top-k tespit oranı", fontsize=9)
    axes[1].tick_params(axis="x", labelsize=7)
    axes[1].tick_params(axis="y", labelsize=7)
    for index, row in classwise.iterrows():
        axes[1].text(index, row["rate"] + 0.03, f"{int(row['detected_count'])}/250", ha="center", fontsize=7)

    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def make_tables(data: dict) -> dict[str, pd.DataFrame]:
    ledger = data["ledger"]

    def row(category: str, run_id: str, score_mode: str | None = None) -> pd.Series:
        subset = ledger[(ledger["category"] == category) & (ledger["run_id"] == run_id)]
        if score_mode is not None:
            subset = subset[subset["score_mode"] == score_mode]
        return subset.iloc[0]

    main_rows = [
        ("AE-MSE 40e", row("training_run", "ae_mse_l128"), "Ara baseline; DRUSEN 88/250"),
        ("AE-MSE 60e", row("training_run", "ae_mse_l128_e60"), "Daha düşük val loss; MSE skorunda benzer"),
        ("AE top-k 60e", row("score_ablation_image_level", "ae_mse_l128_e60", "topk_mse_5"), "Final image-level aday; DRUSEN 95/250"),
        ("VAE+KL", row("training_run", "vae_msekl_l128"), "Beklenen iyileştirmeyi sağlamadı; DRUSEN 41/250"),
        ("AE-L1", row("training_run", "ae_l1_l128"), "Piksel farkı duyarlılığı düşük kaldı"),
        ("AE-MSE+SSIM", row("training_run", "ae_mse_ssim_l128"), "Bu ayarda en zayıf eğitim loss'u"),
        ("Latent 64", row("training_run", "ae_mse_l64"), "Latent ablasyonu; MSE güçlü, top-k geride"),
        ("Latent 256", row("training_run", "ae_mse_l256"), "Latent ablasyonu; stabil fakat finali geçmedi"),
        ("Batch 16", row("training_run", "ae_mse_l256_bs16"), "Batch ablasyonu; top-k güçlü"),
        ("Batch 64", row("training_run", "ae_mse_l256_bs64"), "Batch 16/32'ye göre zayıfladı"),
        ("Content crop", row("training_run", "ae_mse_l256_bs16_crop"), "DRUSEN arttı ama genel performans düştü"),
        ("Retina margin crop", row("training_run", "ae_mse_l256_bs32_retina_margin"), "DRUSEN 96/250; FPR ve AUROC zayıf"),
    ]
    main = pd.DataFrame(
        [
            {
                "Deney": label,
                "AUROC": fmt(series["auroc"]),
                "F1": fmt(series["f1"]),
                "Recall": fmt(series["recall"]),
                "FPR": fmt(series["fpr"]),
                "Kısa yorum": note,
            }
            for label, series, note in main_rows
        ]
    )

    ci = data["bootstrap"]
    best_image = data["best_image"]
    best_patient = data["best_patient"]
    final_summary = pd.DataFrame(
        [
            {
                "Seviye": "Image",
                "Score": "topk_mse_5",
                "AUROC": short_ci(ci, "image", "topk_mse_5", "auroc"),
                "F1": short_ci(ci, "image", "topk_mse_5", "f1"),
                "Recall": fmt(best_image["recall"]),
                "Precision": fmt(best_image["precision"]),
                "FPR": fmt(best_image["fpr"]),
            },
            {
                "Seviye": "Patient mean",
                "Score": "topk_mse_5",
                "AUROC": short_ci(ci, "patient_mean", "topk_mse_5", "auroc"),
                "F1": short_ci(ci, "patient_mean", "topk_mse_5", "f1"),
                "Recall": fmt(best_patient["recall"]),
                "Precision": fmt(best_patient["precision"]),
                "FPR": fmt(best_patient["fpr"]),
            },
        ]
    )

    classwise = data["classwise_score"]
    classwise = classwise[classwise["score_mode"] == "topk_mse_5"].copy()
    classwise["Tespit"] = classwise["detected_count"].astype(int).astype(str) + "/" + classwise["sample_count"].astype(int).astype(str)
    classwise["Oran"] = (classwise["detected_count"] / classwise["sample_count"]).map(lambda value: fmt(value))
    classwise = classwise[["class_name", "mean_score", "Tespit", "Oran"]].rename(
        columns={"class_name": "Sınıf", "mean_score": "Ort. score"}
    )
    classwise["Ort. score"] = classwise["Ort. score"].map(lambda value: fmt(value))

    dataset = data["dataset"].copy()
    dataset["Hasta"] = dataset["patient_count"].astype(int)
    dataset["Görüntü"] = dataset["image_count"].astype(int)
    dataset = dataset[["split_name", "class_name", "Görüntü", "Hasta"]].rename(
        columns={"split_name": "Bölme", "class_name": "Sınıf"}
    )

    literature = pd.DataFrame(
        [
            {
                "Küme": "OCT sınıflandırma",
                "Örnek": "Kermany et al. [1]",
                "Projeye etkisi": "Veri seti ve sınıflar için temel referans.",
            },
            {
                "Küme": "Generative AD",
                "Örnek": "AnoGAN, GANomaly [6], [7]",
                "Projeye etkisi": "Rekonstrüksiyon tabanlı anomaly detection motivasyonu.",
            },
            {
                "Küme": "VAE / SSIM",
                "Örnek": "Kingma, Wang [4], [5]",
                "Projeye etkisi": "VAE+KL ve MSE+SSIM denemeleri için dayanak.",
            },
            {
                "Küme": "Retinal OCT AD",
                "Örnek": "Seebock, Luo, Wang [9], [11], [12]",
                "Projeye etkisi": "DRUSEN zorluğu ve lokal hata analizi için bağlam.",
            },
            {
                "Küme": "Sınırlılık",
                "Örnek": "ProxyAno [10]",
                "Projeye etkisi": "Anomalilerin de iyi reconstruct edilebilme riski.",
            },
        ]
    )

    config = data["config"]
    history = data["history"]
    setup = pd.DataFrame(
        [
            {"Ayar": "Ana model", "Değer": "4 bloklu convolutional autoencoder"},
            {"Ayar": "Final run", "Değer": config["run_id"]},
            {"Ayar": "Giriş", "Değer": f"{config['image_size']}x{config['image_size']} gri seviye, [0,1] normalizasyon"},
            {"Ayar": "Latent / batch", "Değer": f"{config['latent_dim']} / {config['batch_size']}"},
            {"Ayar": "Optimizer", "Değer": f"Adam, lr={config['learning_rate']}"},
            {"Ayar": "Epoch / patience", "Değer": f"{config['epochs']} / {config['early_stopping_patience']}"},
            {"Ayar": "Ana threshold", "Değer": "validation NORMAL dağılımından p95"},
            {"Ayar": "Eğitim süresi", "Değer": f"{history['training_time_sec'] / 60:.1f} dakika, best epoch {history['best_epoch']}"},
            {"Ayar": "Cihaz", "Değer": get_gpu_name()},
        ]
    )

    score_source = pd.read_csv(OUTPUT_DIR / "score_ablation" / "ae_mse_l128_e60" / "score_comparison.csv")
    wanted_scores = [
        "mse",
        "retina_weighted_mse",
        "topk_mse_5",
        "topk_mse_10",
        "ensemble_mse_ssim_topk",
        "ensemble_all_base",
        "l1",
        "mse_ssim_score",
    ]
    score_rows = score_source[score_source["score_mode"].isin(wanted_scores)].copy()
    score_rows["order"] = score_rows["score_mode"].map({name: index for index, name in enumerate(wanted_scores)})
    score_rows = score_rows.sort_values("order")
    score_table = pd.DataFrame(
        [
            {
                "Score": row["score_mode"],
                "AUROC": fmt(row["auroc"]),
                "F1": fmt(row["f1"]),
                "Recall": fmt(row["recall"]),
                "FPR": fmt(row["fpr"]),
            }
            for _, row in score_rows.iterrows()
        ]
    )

    threshold_source = pd.read_csv(OUTPUT_DIR / "score_ablation" / "ae_mse_l128_e60" / "threshold_comparison.csv")
    threshold_rows = threshold_source[threshold_source["score_mode"] == "topk_mse_5"].copy()
    threshold_table = pd.DataFrame(
        [
            {
                "Eşik": f"p{int(row['percentile'])}",
                "Threshold": fmt(row["threshold"], 5),
                "Precision": fmt(row["precision"]),
                "Recall": fmt(row["recall"]),
                "F1": fmt(row["f1"]),
                "FPR": fmt(row["fpr"]),
            }
            for _, row in threshold_rows.iterrows()
        ]
    )

    return {
        "literature": literature,
        "setup": setup,
        "main": main,
        "score": score_table,
        "threshold": threshold_table,
        "final": final_summary,
        "classwise": classwise,
        "dataset": dataset,
    }


def build_markdown(data: dict, tables: dict[str, pd.DataFrame]) -> str:
    config = data["config"]
    history = data["history"]
    best_image = data["best_image"]
    best_patient = data["best_patient"]
    training_minutes = history["training_time_sec"] / 60.0
    best_epoch = history["best_epoch"]
    gpu = get_gpu_name()

    sections = report_sections(data, tables, gpu)
    lines = [f"# {TITLE_TR}", "", f"**English title:** {TITLE_EN}", ""]
    lines.append("## Özet")
    lines.append(sections["Özet"])
    lines.append("")
    lines.append("## Abstract")
    lines.append(sections["Abstract"])
    lines.append("")
    lines.append("## Anahtar Kelimeler / Keywords")
    lines.append(sections["Anahtar Kelimeler / Keywords"])
    lines.append("")
    for heading in ["Giriş", "Literatür", "Yöntem", "Bulgular", "Sonuç", "Kaynakça"]:
        lines.append(f"## {heading}")
        if heading == "Yöntem":
            lines.append(sections[heading])
            lines.append("")
            lines.append(markdown_table(tables["dataset"]))
            lines.append("")
            lines.append(markdown_table(tables["setup"]))
            lines.append("")
            lines.append("![Şekil 1. Önerilen final deney hattı.](../outputs/final_report/final_pipeline.png)")
        elif heading == "Literatür":
            lines.append(sections[heading])
            lines.append("")
            lines.append(markdown_table(tables["literature"]))
        elif heading == "Bulgular":
            lines.append(sections[heading])
            lines.append("")
            lines.append(markdown_table(tables["main"]))
            lines.append("")
            lines.append(markdown_table(tables["score"]))
            lines.append("")
            lines.append(markdown_table(tables["final"]))
            lines.append("")
            lines.append(markdown_table(tables["threshold"]))
            lines.append("")
            lines.append(markdown_table(tables["classwise"]))
            lines.append("")
            lines.append(FIGURE_EXPLANATION)
            lines.append("")
            lines.append("![Şekil 2. Final performans ve sınıf bazlı tespit özeti.](../outputs/final_report/final_results_summary.png)")
            lines.append("")
            lines.append("![Şekil 3. Top-k residual nitel analiz örnekleri.](../outputs/score_ablation/ae_mse_l128_e60/topk_explainability_grid.png)")
            lines.append("")
            lines.append("![Şekil 4. DRUSEN false negative örnekleri.](../outputs/score_ablation/ae_mse_l128_e60/drusen_false_negative_topk_grid.png)")
        elif heading == "Kaynakça":
            lines.extend(REFERENCES)
        else:
            lines.append(sections[heading])
        lines.append("")

    return "\n".join(lines)


def get_gpu_name() -> str:
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
            check=True,
            capture_output=True,
            text=True,
            timeout=10,
        )
    except Exception:
        return "CUDA destekli yerel ortam"
    return result.stdout.strip().splitlines()[0] if result.stdout.strip() else "CUDA destekli yerel ortam"


def report_sections(data: dict, tables: dict[str, pd.DataFrame], gpu: str) -> dict[str, str]:
    config = data["config"]
    history = data["history"]
    best_image = data["best_image"]
    best_patient = data["best_patient"]
    train_normal = int(data["dataset"][(data["dataset"]["split_name"] == "train") & (data["dataset"]["class_name"] == "NORMAL")]["image_count"].iloc[0])
    val_normal = int(data["dataset"][(data["dataset"]["split_name"] == "val") & (data["dataset"]["class_name"] == "NORMAL")]["image_count"].iloc[0])
    training_minutes = history["training_time_sec"] / 60.0

    return {
        "Özet": (
            "Bu çalışmada retinal OCT görüntülerinde patolojik örneklerin tespiti için yalnızca normal görüntülerle eğitilen "
            "konvolüsyonel autoencoder tabanlı bir anomali tespit hattı geliştirilmiştir. Kermany OCT2017/Mendeley veri kümesinde "
            "eğitim aşamasına sadece NORMAL sınıfı alınmış, doğrulama bölmesi hasta kimliği kullanılarak ayrılmış ve test aşamasında "
            "CNV, DME, DRUSEN ve NORMAL örnekleri binary anomaly detection problemi olarak değerlendirilmiştir. Final aşamasında ara "
            "baseline korunmuş; VAE+KL, L1, MSE+SSIM, latent boyut, batch size, crop ön işleme, farklı anomaly score'lar, hasta düzeyi "
            "birleştirme ve bootstrap güven aralığı denemeleri eklenmiştir. En iyi image-level sonuç, 60 epoch eğitilen AE-MSE modelinin "
            "topk_mse_5 score'u ile AUROC "
            f"{best_image['auroc']:.4f}, F1 {best_image['f1']:.4f}, precision {best_image['precision']:.4f}, recall {best_image['recall']:.4f} "
            f"ve FPR {best_image['fpr']:.4f} olarak bulunmuştur. Hasta düzeyinde mean aggregation ile F1 {best_patient['f1']:.4f} elde edilmiştir. "
            "Sonuçlar top-k residual score'un klasik ortalama MSE'ye göre daha ayırt edici olduğunu, buna karşılık DRUSEN örneklerinin hala en zor patolojik grup olduğunu göstermektedir."
        ),
        "Abstract": (
            "This project develops a convolutional autoencoder pipeline for detecting pathological retinal OCT images by learning only from normal samples. "
            "Using the Kermany OCT2017/Mendeley dataset, only NORMAL images were used for training, the validation subset was separated at patient level, "
            "and CNV, DME, DRUSEN and NORMAL test images were evaluated as a binary anomaly detection task. In the final stage, the baseline was extended with "
            "VAE+KL, L1, MSE+SSIM, latent-size and batch-size ablations, crop preprocessing trials, alternative anomaly scores, patient-level aggregation and bootstrap confidence intervals. "
            f"The best image-level setting was the 60-epoch AE-MSE model with topk_mse_5 scoring, reaching AUROC {best_image['auroc']:.4f}, F1 {best_image['f1']:.4f}, "
            f"precision {best_image['precision']:.4f}, recall {best_image['recall']:.4f} and FPR {best_image['fpr']:.4f}. Patient-level mean aggregation achieved F1 {best_patient['f1']:.4f}. "
            "The findings indicate that localized top-k residual scoring improves separability over mean MSE, while DRUSEN remains the most challenging pathology group."
        ),
        "Anahtar Kelimeler / Keywords": (
            "retinal OCT, anomaly detection, autoencoder, reconstruction error, top-k residual, VAE, SSIM, medical imaging, deep learning"
        ),
        "Giriş": (
            "Optik koherens tomografi (OCT), retina katmanlarını yüksek çözünürlükte görüntüleyebildiği için CNV, DME ve DRUSEN gibi patolojilerin değerlendirilmesinde önemli bir görüntüleme yöntemidir. "
            "Buna karşın gerçek klinik senaryolarda tüm patoloji türleri için yeterli ve güvenilir etiket toplamak her zaman kolay değildir. Denetimli sınıflandırma modelleri güçlü sonuçlar verebilse de, "
            "etiket bağımlılığı ve dağılım dışı örneklere karşı kırılganlık bu tür sistemlerin pratik kullanımında önemli sınırlardır [1], [3]. Bu nedenle bu projede hastalık türünü doğrudan sınıflandırmak yerine, "
            "normal retina anatomisini öğrenen bir modelin patolojik görüntülerde daha yüksek rekonstrüksiyon hatası üretmesi fikri incelenmiştir."
            "\n\n"
            "Ara rapor aşamasında çalışan bir AE-MSE baseline kurulmuştu. Final aşamasında amaç sadece tek bir sonucu iyileştirmek değil, aynı zamanda hangi teknik kararların işe yaradığını sistematik olarak görmekti. "
            "Bu kapsamda VAE tabanlı üretken latent model, L1 ve SSIM içeren loss seçenekleri, latent boyut ve batch size ablasyonları, boş alanları azaltmaya yönelik crop denemeleri, top-k residual score, hasta düzeyi değerlendirme, "
            "bootstrap güven aralığı ve nitel heatmap görselleri aynı deney hattına eklenmiştir. Bu nedenle final raporu ara rapordaki başlangıç fikrini temel almakla birlikte, asıl olarak finalde yapılan genişletmeleri ve bunların kontrollü karşılaştırmasını sunmaktadır. "
            "Böylece proje, yalnızca bir model çalıştırma örneği olmaktan çıkarılıp savunulabilir bir karşılaştırmalı analiz haline getirilmiştir."
        ),
        "Literatür": (
            "Kermany ve arkadaşlarının OCT2017 veri kümesi ve derin öğrenme çalışması, retina OCT görüntülerinde denetimli tanı sınıflandırmasının güçlü bir örneğini sunar [1], [2]. Daha genel tıbbi görüntüleme literatürü de derin öğrenmenin yüksek performans potansiyelini göstermektedir [3]. "
            "Ancak bu çizgideki modeller çoğunlukla her sınıf için etiketli örnek gerektirir. Anomali tespiti literatüründe AnoGAN, GANomaly ve DRAEM gibi çalışmalar, normal dağılımdan sapmayı rekonstrüksiyon veya adversarial öğrenme yoluyla yakalamaya çalışmıştır [6]-[8]. "
            "VAE yaklaşımı, olasılıksal latent uzay yapısıyla bu probleme doğal bir alternatif sunar [4]. SSIM ise piksel farkından farklı olarak yapısal benzerliği ölçtüğü için rekonstrüksiyon kalitesini değerlendirmede anlamlı bir bileşen olabilir [5]."
            "\n\n"
            "Retinal OCT özelinde Seebock ve arkadaşları belirsizlik tabanlı anomali tespiti, Luo ve arkadaşları çok çözünürlüklü autoencoder, Wang ve arkadaşları ise zayıf denetimli anomali segmentasyonu üzerine çalışmıştır [9], [11], [12]. "
            "ProxyAno çalışması ise rekonstrüksiyon tabanlı yöntemlerin önemli bir riskini vurgular: model bazı anomalileri de iyi reconstruct edebilir ve bu durum duyarlılığı düşürebilir [10]. Bu proje bu literatürü doğrudan büyük bir modelle yeniden üretmek yerine, OCT2017 üzerinde normal-only, hasta düzeyinde kontrollü ve tekrar üretilebilir bir baseline ile bu riskleri deneysel olarak incelemektedir."
            "\n\n"
            "Literatürdeki bu çalışmaların önemli bir kısmı daha karmaşık mimariler, ek anotasyonlar veya özel ön işleme varsayımları kullanmaktadır. Bu projede bilinçli olarak daha sade bir mimariyle başlanmıştır; çünkü final hedefi yalnızca yüksek bir tek skor üretmek değil, hangi kararın hangi sonucu değiştirdiğini açıklayabilmektir. Bu nedenle VAE, SSIM, crop ve score ensemble gibi seçenekler aynı veri protokolü altında karşılaştırılmıştır. Böyle bir kurgu, daha gelişmiş yöntemlere geçmeden önce güvenilir bir baseline oluşturmayı ve DRUSEN gibi zor sınıflarda hatanın nereden geldiğini tartışmayı kolaylaştırmaktadır."
        ),
        "Yöntem": (
            f"Veri seti olarak Kermany OCT2017/Mendeley V3 kullanılmıştır. Klasör sözleşmesi data/oct2017/{{train,test}}/{{NORMAL,CNV,DME,DRUSEN}} biçiminde tutulmuştur. Eğitimde {train_normal} NORMAL görüntü, doğrulamada {val_normal} NORMAL görüntü kullanılmıştır; test kümesi ise her sınıftan 250 görüntü olmak üzere toplam 1000 görüntü içermektedir. "
            "Dosya adlarındaki hasta kimliği parse edilerek train ve validation arasında hasta kesişimi engellenmiştir. Patient-level aggregation sırasında farklı sınıflardaki aynı sayısal ID'lerin karışmaması için class prefix korunmuştur; örneğin NORMAL-0101 ve CNV-0101 farklı hasta anahtarları olarak tutulmuştur. Eğitim verisine hiçbir patolojik sınıf alınmamış, eşik seçimi yalnızca validation NORMAL score dağılımından yapılmıştır. Tüm görüntüler gri seviyeye çevrilmiş, 128x128 boyutuna getirilmiş ve [0,1] aralığında normalize edilmiştir."
            "\n\n"
            "Ana model dört bloklu konvolüsyonel autoencoder yapısındadır. Encoder, görüntüyü düşük boyutlu latent temsile indirger; decoder aynı uzaydan görüntüyü yeniden oluşturur. Ana koşuda latent_dim=128, Adam optimizer, learning rate 1e-3, batch size 32, maksimum 60 epoch ve patience 10 kullanılmıştır. Eğitim NVIDIA GeForce RTX 4060 Laptop GPU üzerinde yapılmış ve final koşu yaklaşık "
            f"{training_minutes:.1f} dakika sürmüştür. VAE denemesinde encoder mu ve logvar üretmiş, reparameterization trick kullanılmış ve loss MSE + 1e-4 * KL olarak alınmıştır. SSIM denemesinde loss 0.8*MSE + 0.2*(1-SSIM), L1 denemesinde ise ortalama mutlak hata kullanılmıştır."
            "\n\n"
            "Deney tasarımında iki ayrım korunmuştur. Birincisi, training loss ve anomaly score her zaman aynı olmak zorunda değildir; örneğin model MSE ile eğitildikten sonra aynı rekonstrüksiyonlardan farklı score türleri hesaplanmıştır. İkincisi, eşik seçimi ile performans değerlendirmesi ayrılmıştır; p95 eşiği validation normal skorlarından alınmış, test verisi yalnızca bu eşik sabitlendikten sonra değerlendirilmiştir. Bu ayrım, test setine göre threshold ayarlama riskini azaltmıştır."
            "\n\n"
            "Başlangıç anomaly score'u piksel başına ortalama MSE'dir. Finalde daha lokal patolojik farklılıkları yakalamak için topk_mse_5 score'u kullanılmıştır; bu score, kare hata haritasındaki en yüksek yüzde 5 pikselin ortalamasını alır. Score ablation aşamasında mse, l1, ssim_error, mse_ssim_score, retina_band_mse, retina_weighted_mse, topk_mse_5, topk_mse_10, ensemble_mse_l1, ensemble_mse_l1_ssim, ensemble_mse_retina, ensemble_mse_ssim_topk, ensemble_l1_ssim_topk, ensemble_mse_l1_retina ve ensemble_all_base seçenekleri hesaplanmıştır. Tablo V, bu geniş listenin rapor içinde okunabilir kalması için en açıklayıcı alt kümesini göstermektedir; tam CSV çıktıları proje klasöründe saklanmıştır. Ana operasyon eşiği tüm modeller için p95 olarak tutulmuş, p97 ve p99 değerleri yalnızca karşılaştırma amacıyla saklanmıştır. Bu tercih önemlidir; çünkü test seti threshold seçiminde kullanılmamış, test yalnızca son değerlendirme için ayrılmıştır. Hasta düzeyi değerlendirmede aynı hasta ID'sine ait görüntü score'ları mean aggregation ile birleştirilmiştir. Sonuçların belirsizliğini görmek için bootstrap ile AUROC ve F1 için yüzde 95 güven aralıkları hesaplanmıştır. Nitel analiz için residual heatmap overlay, top-k residual grid ve DRUSEN false negative örnekleri üretilmiştir. Bu görseller pixel-level anotasyon olmadığı için segmentasyon çıktısı olarak değil, modelin hangi bölgelerde yüksek hata ürettiğini gösteren açıklayıcı destek olarak yorumlanmıştır."
            "\n\n"
            "Metrik tarafında AUROC, threshold'dan bağımsız genel ayrıştırma gücünü verdiği için ana metrik olarak seçilmiştir. Precision, patolojik olarak işaretlenen görüntülerin ne kadarının gerçekten patolojik olduğunu; recall, patolojik örneklerin ne kadarının yakalandığını; F1 ise precision-recall dengesini göstermektedir. FPR özellikle önemlidir, çünkü normal görüntüleri gereksiz yere patolojik işaretlemek gerçek klinik akışta ek inceleme yükü oluşturabilir. Bu nedenle final model seçimi yalnızca AUROC değerine göre değil, F1, recall, precision ve FPR birlikte değerlendirilerek yapılmıştır."
            "\n\n"
            "Ön işleme denemeleri ayrı bir karar noktası olarak ele alınmıştır. OCT görüntülerinde üst veya alt bölgede siyah/beyaz boş alanlar bulunabildiği için content crop ve retina-margin crop denenmiştir. Ancak bu kesme işlemleri bazı örneklerde retina çevresindeki bağlamı değiştirme riski taşımaktadır. Bu yüzden crop yaklaşımları doğrudan final modele alınmamış, önce önizleme görselleriyle incelenmiş, ardından yalnızca anlamlı görünen seçenekler tam eğitim koşusuna taşınmıştır. Sonuçlar, görsel olarak makul görünen crop işlemlerinin bile anomaly score dağılımını her zaman iyileştirmediğini göstermiştir."
        ),
        "Bulgular": (
            "Tablo IV final aşamasında yapılan ana deneyleri özetlemektedir. Ara baseline olan AE-MSE 40 epoch koşusu AUROC 0.9104 ve F1 0.8262 üretmiştir. Aynı model 60 epoch eğitildiğinde validation loss düşmüş, ancak klasik ortalama MSE score'u tek başına büyük bir sıçrama sağlamamıştır. Asıl iyileşme score tasarımından gelmiştir: 60 epoch AE-MSE modelinde topk_mse_5 kullanıldığında AUROC 0.9457 ve F1 0.8464'e çıkmıştır. Bu sonuç, OCT patolojilerinin tüm görüntüye yayılmak yerine sınırlı bölgelerde yoğunlaşabildiğini ve ortalama MSE'nin bu sinyali zayıflatabildiğini göstermektedir."
            "\n\n"
            "VAE+KL, L1 ve MSE+SSIM denemeleri beklenen iyileştirmeyi sağlamamıştır. VAE özellikle DRUSEN sınıfında zayıf kalmış, MSE+SSIM ise bu veri ve mimari ayarında en düşük genel performansı vermiştir. Latent 64/128/256 karşılaştırmasında değerler birbirine yakın kalmış, final aday için latent 128 korunmuştur. Batch size 16 güçlü sonuç vermesine rağmen final score'da 60 epoch latent 128 koşusunu geçememiştir. Crop denemelerinde content crop ve retina-margin crop DRUSEN yakalama sayısını bir miktar artırmış, fakat AUROC ve FPR tarafında genel performansı düşürmüştür; bu nedenle final modelde crop kullanılmamıştır."
            "\n\n"
            "Score ablation sonuçları, iyileştirmenin model mimarisinden çok hata haritasının nasıl özetlendiğiyle ilişkili olduğunu göstermiştir. Tablo V bu farkı sayısal olarak göstermektedir: topk_mse_5, ortalama MSE'ye göre AUROC değerini 0.9112'den 0.9457'ye çıkarmıştır. Retina-band ve retina-weighted MSE skorları ortalama MSE'ye yakın kalmış, L1/SSIM tabanlı ensemble score'lar ise topk_mse_5'i geçememiştir. Border crop yalnızca önizleme düzeyinde bırakılmıştır; çünkü ilk görsel inceleme bu yaklaşımın sınırlı katkı sağlayacağını göstermiştir. Content crop ve retina-margin crop ise tam eğitimle denenmiş, ancak daha yüksek DRUSEN tespitine rağmen genel ayrımı bozduğu için final seçime alınmamıştır."
            "\n\n"
            "Tablo VI final image-level ve patient-level sonuçlarını bootstrap güven aralıklarıyla birlikte vermektedir. Image-level p95 eşiğinde precision 0.9911 ve FPR 0.0200 elde edilmiştir; yani normal test görüntülerinde yanlış pozitif sayısı 5/250 seviyesinde kalmıştır. Hasta düzeyinde mean aggregation "
            f"AUROC {best_patient['auroc']:.4f}, F1 {best_patient['f1']:.4f}, recall {best_patient['recall']:.4f}, precision {best_patient['precision']:.4f} ve FPR {best_patient['fpr']:.4f} üretmiştir. "
            "Tablo VII, p95-p97-p99 eşiğinin precision-recall dengesini nasıl değiştirdiğini göstermektedir. p97 ve p99 precision değerini korusa da recall belirgin biçimde düştüğü için final operasyon noktası p95 olarak bırakılmıştır. Tablo VIII, sınıf bazlı davranışı göstermektedir. CNV 247/250, DME 212/250 yakalanırken DRUSEN 95/250 seviyesinde kalmıştır. Bu durum DRUSEN'in normal anatomiden daha ince ve lokal sapmalar içermesiyle uyumludur. Üretilen heatmap ve DRUSEN false negative görselleri de modelin bazı DRUSEN örneklerinde patolojik bölgeyi yeterince yüksek score'a taşıyamadığını göstermiştir."
            "\n\n"
            "60 epoch denemesi ayrıca eğitim süresi ve overfitting açısından da yorumlanmıştır. En iyi validation loss 58. epoch'ta görülmüş, bu da 40 epoch sınırının biraz erken kalabileceğini göstermiştir. Buna rağmen ortalama MSE score'u çok az değişmiştir; dolayısıyla ek epoch tek başına yeterli olmamıştır. Final iyileştirme, daha uzun eğitim ile top-k residual score'un birlikte kullanılmasından gelmiştir. Bu gözlem, rekonstrüksiyon tabanlı anomaly detection problemlerinde yalnızca loss düşüşünü izlemek yerine, score fonksiyonunun patolojik sinyali nasıl özetlediğini de ayrıca test etmek gerektiğini göstermektedir."
            "\n\n"
            "VAE sonucunun zayıf kalması da önemli bir bulgudur. VAE'nin daha düzenli bir latent uzay oluşturması beklenebilir; ancak KL terimi rekonstrüksiyonu daha pürüzsüz hale getirerek küçük ve lokal patolojik farklılıkların score'a yansımasını azaltmış olabilir. Benzer şekilde SSIM içeren loss yapısal benzerliği gözetse de, bu veri ve 128x128 çözünürlükte patoloji-normal ayrımını güçlendirmemiştir. Bu sonuçlar negatif görünse de final rapor açısından değerlidir; çünkü hangi geliştirmelerin gerçekten katkı verdiğini, hangilerinin yalnızca teorik olarak cazip kaldığını ayırmaktadır."
        ),
        "Sonuç": (
            "Bu projede normal retina OCT görüntülerinden öğrenen konvolüsyonel autoencoder tabanlı bir anomali tespit sistemi uçtan uca geliştirilmiş ve final aşamasında kapsamlı ablasyonlarla değerlendirilmiştir. En güçlü sonuç, yeni ve daha karmaşık bir modelden değil, aynı AE-MSE modelinin lokal hataya duyarlı top-k residual score ile değerlendirilmesinden gelmiştir. "
            f"Final image-level sonuç AUROC {best_image['auroc']:.4f} ve F1 {best_image['f1']:.4f}, hasta düzeyi sonuç ise AUROC {best_patient['auroc']:.4f} ve F1 {best_patient['f1']:.4f}'dır."
            "\n\n"
            "Deneyler VAE, L1, MSE+SSIM, latent boyut, batch size ve crop gibi seçeneklerin bu veri ve ayarlarda her zaman iyileştirme getirmediğini göstermiştir. Özellikle crop işlemleri DRUSEN yakalamayı artırsa bile genel hata dengesini bozmuştur. Çalışmanın ana sınırlılıkları, test setinin sınıf başına 250 görüntüyle sınırlı tutulması, pixel-level lezyon anotasyonu bulunmaması ve modelin klinik karar destek sistemi olarak kullanılmaya hazır olmamasıdır. Buna rağmen hasta bazlı split, testten bağımsız threshold seçimi, bootstrap güven aralığı ve nitel residual analizleriyle proje, final teslimi için tutarlı ve tekrar üretilebilir bir normal-only OCT anomaly detection baseline sunmaktadır."
            "\n\n"
            "Final sonrası en doğal geliştirme yönleri, daha yüksek çözünürlükte eğitim, dış veri setinde doğrulama, DRUSEN odaklı ek analiz ve pixel-level anotasyon varsa gerçek anomaly localization değerlendirmesidir. Bununla birlikte bu çalışma kapsamında hedeflenen ana çıktı tamamlanmıştır: normal-only öğrenme yaklaşımı çalıştırılmış, başarısız denemeler saklanmadan raporlanmış ve final model seçimi yalnızca test performansına göre değil, validation tabanlı threshold, sınıf bazlı davranış ve hasta düzeyi analiz birlikte düşünülerek yapılmıştır."
            "\n\n"
            "Sonuçların aktarılabilirliği konusunda dikkatli olunmalıdır. Bu çalışmada test seti sınıf başına 250 görüntüyle dengeli tutulmuştur; gerçek klinik akışta patolojik ve normal örnek oranları farklı olabilir. Bu durumda aynı precision, recall ve FPR değerleri doğrudan aynı kullanım etkisine karşılık gelmeyebilir. Ayrıca OCT2017 görüntüleri belirli cihaz ve veri toplama koşullarından geldiği için farklı merkezlerden gelen görüntülerde dış doğrulama yapılmadan klinik genelleme iddiasında bulunmak doğru değildir."
            "\n\n"
            "Buna rağmen proje, final ders kapsamı açısından üretken yapay zeka ve derin öğrenme yöntemlerinin tıbbi görüntülerde nasıl kontrollü deney tasarımına dönüştürülebileceğini göstermektedir. VAE gibi üretken latent model denenmiş, ancak sonuçlar iyi çıkmadığında saklanmamıştır. SSIM, L1, crop ve ensemble score denemeleri de aynı şekilde rapora dahil edilmiştir. Bu yaklaşım, yalnızca başarılı sonucu sunmak yerine model geliştirme sürecinin karar noktalarını görünür kıldığı için final projesinin teknik değerini artırmaktadır."
        ),
    }


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_element = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(cols_element)
    cols_element.set(qn("w:num"), str(count))


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for element in list(body):
        if not element.tag.endswith("sectPr"):
            body.remove(element)


def set_table_borders(table, color: str = "000000") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_margins(table, margin_twips: int = 45) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin in ("top", "left", "bottom", "right"):
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def configure_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(9.8)
    for section in doc.sections:
        section.top_margin = Inches(0.70)
        section.bottom_margin = Inches(0.70)
        section.left_margin = Inches(0.66)
        section.right_margin = Inches(0.66)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(TITLE_TR)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(TITLE_EN)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13.2)

    table = doc.add_table(rows=1, cols=len(AUTHORS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_none(table)
    for index, author in enumerate(AUTHORS):
        cell = table.rows[0].cells[index]
        lines = [
            author["name"],
            author["department"],
            author["institution"],
            author["location"],
            author["email"],
        ]
        for line_index, line in enumerate(lines):
            p = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            if line_index == 0:
                run.bold = True


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{label} - ")
    run.bold = True
    run.italic = True
    p.add_run(text)


def add_heading(doc: Document, number: int | None, heading: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    text = f"{roman(number)}. {tr_upper(heading)}" if number is not None else tr_upper(heading)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.0)


def roman(number: int) -> str:
    values = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = ""
    remaining = number
    for value, symbol in values:
        while remaining >= value:
            result += symbol
            remaining -= value
    return result


def tr_upper(text: str) -> str:
    translation = str.maketrans({"i": "İ", "ı": "I", "ğ": "Ğ", "ü": "Ü", "ş": "Ş", "ö": "Ö", "ç": "Ç"})
    return text.translate(translation).upper()


def add_body_text(doc: Document, text: str) -> None:
    for paragraph_text in text.split("\n\n"):
        p = doc.add_paragraph(paragraph_text.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0


def add_dataframe_table(doc: Document, frame: pd.DataFrame, caption: str, font_size: float = 6.2) -> None:
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(7.5)
        run.bold = True

    table = doc.add_table(rows=1 + len(frame), cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    set_table_margins(table)
    for col_idx, header in enumerate(frame.columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)

    for row_idx, row in enumerate(frame.itertuples(index=False, name=None), start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)


def add_figure(doc: Document, path: Path, caption: str, width: float = 3.0, height: float | None = None) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    if height is None:
        p.add_run().add_picture(str(path), width=Inches(width))
    else:
        p.add_run().add_picture(str(path), height=Inches(height))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(7.3)


def add_column_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.COLUMN)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_references(doc: Document) -> None:
    add_heading(doc, None, "Kaynakça")
    for reference in REFERENCES:
        p = doc.add_paragraph(reference)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(7.8)


def build_docx(data: dict, tables: dict[str, pd.DataFrame], docx_path: Path) -> None:
    sections = report_sections(data, tables, get_gpu_name())
    doc = Document(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else Document()
    clear_document(doc)
    configure_styles(doc)
    set_columns(doc.sections[0], 1)
    add_title_block(doc)

    body_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_columns(body_section, 2)

    add_labeled_paragraph(doc, "Özet", sections["Özet"])
    add_labeled_paragraph(doc, "Abstract", sections["Abstract"])
    add_labeled_paragraph(doc, "Anahtar Kelimeler / Keywords", sections["Anahtar Kelimeler / Keywords"])

    add_heading(doc, 1, "Giriş")
    add_body_text(doc, sections["Giriş"])
    add_heading(doc, 2, "Literatür")
    add_body_text(doc, sections["Literatür"])
    add_dataframe_table(doc, tables["literature"], "Tablo I. Literatürdeki ana kümeler ve bu projedeki karşılıkları.", font_size=6.2)
    add_heading(doc, 3, "Yöntem")
    add_body_text(doc, sections["Yöntem"])
    add_dataframe_table(doc, tables["dataset"], "Tablo II. Kullanılan veri bölmeleri ve hasta sayıları.", font_size=6.5)
    add_dataframe_table(doc, tables["setup"], "Tablo III. Final ana koşu ve eğitim ayarları.", font_size=6.5)
    add_figure(doc, FINAL_FIGURE_DIR / "final_pipeline.png", "Şekil 1. Önerilen final deney hattı.", width=3.05)

    add_heading(doc, 4, "Bulgular")
    add_body_text(doc, sections["Bulgular"])
    add_column_break(doc)
    add_dataframe_table(doc, tables["main"], "Tablo IV. Final aşamasında yapılan ana deneyler.", font_size=6.1)
    add_dataframe_table(doc, tables["score"], "Tablo V. Final run için score ablation özeti.", font_size=6.2)
    add_dataframe_table(doc, tables["final"], "Tablo VI. Final aday sonuçları ve bootstrap güven aralıkları.", font_size=6.2)
    add_dataframe_table(doc, tables["threshold"], "Tablo VII. topk_mse_5 için eşik duyarlılığı.", font_size=6.3)
    add_column_break(doc)
    add_dataframe_table(doc, tables["classwise"], "Tablo VIII. Final top-k score için sınıf bazlı tespit özeti.", font_size=6.5)
    add_body_text(doc, FIGURE_EXPLANATION)
    add_figure(doc, FINAL_FIGURE_DIR / "final_results_summary.png", "Şekil 2. Final performans ve sınıf bazlı tespit özeti.", width=3.05)
    add_figure(
        doc,
        OUTPUT_DIR / "score_ablation" / "ae_mse_l128_e60" / "topk_explainability_grid.png",
        "Şekil 3. Top-k residual overlay ile nitel analiz örnekleri.",
        width=3.05,
    )
    add_figure(
        doc,
        OUTPUT_DIR / "score_ablation" / "ae_mse_l128_e60" / "drusen_false_negative_topk_grid.png",
        "Şekil 4. DRUSEN false negative örnekleri.",
        width=3.05,
    )

    add_heading(doc, 5, "Sonuç")
    add_body_text(doc, sections["Sonuç"])
    add_references(doc)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    data = load_report_data()
    tables = make_tables(data)
    create_pipeline_figure(FINAL_FIGURE_DIR / "final_pipeline.png")
    create_results_figure(data, FINAL_FIGURE_DIR / "final_results_summary.png")

    markdown = build_markdown(data, tables)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    (REPORT_DIR / "final_rapor.md").write_text(markdown, encoding="utf-8")
    build_docx(data, tables, REPORT_DIR / "Grup12_KorayÖztürk_EmirAlpİlhan_Final_Rapor.docx")


if __name__ == "__main__":
    main()
